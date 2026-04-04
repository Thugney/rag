import re
import logging
import csv
import json
import shutil
import subprocess
import tempfile
import zipfile
from collections import defaultdict
from dataclasses import dataclass, field
from typing import Iterator, List, Dict, Optional, Any, Tuple, Protocol, Sequence
from pathlib import Path
from datetime import datetime
from xml.etree import ElementTree as ET

import PyPDF2
from docx import Document
import numpy as np

logger = logging.getLogger(__name__)

@dataclass
class DocumentChunk:
    """Represents a document chunk with metadata"""
    content: str
    metadata: Dict[str, Any]
    embedding: Optional[np.ndarray] = field(default=None, repr=False)


@dataclass
class ParsedDocumentElement:
    """Represents a parsed structural element before chunking."""
    content: str
    metadata: Dict[str, Any]
    content_type: str = "text"
    structural_role: str = "body"
    parser_name: str = "unknown"
    element_index: int = 0


@dataclass
class ParsedDocument:
    """Represents a parsed document with structured elements."""
    source: str
    filename: str
    elements: List[ParsedDocumentElement]


class WindowsOCREngine:
    """Small wrapper around the built-in Windows OCR API via powershell.exe."""

    _script_path: Optional[Path] = None
    _available: Optional[bool] = None

    def __init__(self, timeout_seconds: int = 60):
        self.timeout_seconds = timeout_seconds

    @classmethod
    def is_available(cls) -> bool:
        if cls._available is not None:
            return cls._available

        if not shutil.which("powershell.exe"):
            cls._available = False
            return cls._available

        command = (
            "Add-Type -AssemblyName System.Runtime.WindowsRuntime; "
            "try { "
            "$null = [Windows.Media.Ocr.OcrEngine, Windows.Foundation, ContentType=WindowsRuntime]; "
            "Write-Output 'available' "
            "} catch { Write-Output 'missing' }"
        )
        try:
            result = subprocess.run(
                ["powershell.exe", "-NoProfile", "-Command", command],
                capture_output=True,
                text=True,
                timeout=20,
                check=False,
            )
            cls._available = "available" in (result.stdout or "").lower()
        except Exception:
            cls._available = False

        return cls._available

    @classmethod
    def _ensure_script_path(cls) -> Path:
        script_path = Path(tempfile.gettempdir()) / "rag-windows-ocr.ps1"
        script_path.write_text(
            """
Add-Type -AssemblyName System.Runtime.WindowsRuntime
function AwaitTyped($op, [Type]$resultType) {
  $method = [System.WindowsRuntimeSystemExtensions].GetMethods() | Where-Object { $_.Name -eq 'AsTask' -and $_.IsGenericMethod -and $_.GetParameters().Count -eq 1 } | Select-Object -First 1
  $generic = $method.MakeGenericMethod($resultType)
  $task = $generic.Invoke($null, @($op))
  return $task.GetAwaiter().GetResult()
}
$null = [Windows.Storage.StorageFile, Windows.Storage, ContentType=WindowsRuntime]
$null = [Windows.Graphics.Imaging.BitmapDecoder, Windows.Foundation, ContentType=WindowsRuntime]
$null = [Windows.Graphics.Imaging.SoftwareBitmap, Windows.Foundation, ContentType=WindowsRuntime]
$null = [Windows.Media.Ocr.OcrEngine, Windows.Foundation, ContentType=WindowsRuntime]
$null = [Windows.Media.Ocr.OcrResult, Windows.Foundation, ContentType=WindowsRuntime]
$null = [Windows.Storage.Streams.IRandomAccessStream, Windows.Storage.Streams, ContentType=WindowsRuntime]
$null = [Windows.Globalization.Language, Windows.Foundation, ContentType=WindowsRuntime]
try {
  $path = $args[0]
  $file = AwaitTyped ([Windows.Storage.StorageFile]::GetFileFromPathAsync($path)) ([Windows.Storage.StorageFile])
  $stream = AwaitTyped ($file.OpenAsync([Windows.Storage.FileAccessMode]::Read)) ([Windows.Storage.Streams.IRandomAccessStream])
  $decoder = AwaitTyped ([Windows.Graphics.Imaging.BitmapDecoder]::CreateAsync($stream)) ([Windows.Graphics.Imaging.BitmapDecoder])
  $softwareBitmap = AwaitTyped ($decoder.GetSoftwareBitmapAsync()) ([Windows.Graphics.Imaging.SoftwareBitmap])
  $engine = [Windows.Media.Ocr.OcrEngine]::TryCreateFromUserProfileLanguages()
  if ($null -eq $engine) {
    $engine = [Windows.Media.Ocr.OcrEngine]::TryCreateFromLanguage([Windows.Globalization.Language]::new('en-US'))
  }
  if ($null -eq $engine) {
    throw 'No Windows OCR engine available.'
  }
  $result = AwaitTyped ($engine.RecognizeAsync($softwareBitmap)) ([Windows.Media.Ocr.OcrResult])
  $lineCount = @($result.Lines).Count
  $wordCount = (@($result.Lines) | ForEach-Object { @($_.Words).Count } | Measure-Object -Sum).Sum
  if ($null -eq $wordCount) { $wordCount = 0 }
  [ordered]@{
    text = $result.Text
    line_count = $lineCount
    word_count = $wordCount
    language = $engine.RecognizerLanguage.LanguageTag
    image_width = $softwareBitmap.PixelWidth
    image_height = $softwareBitmap.PixelHeight
    ocr_confidence = $null
  } | ConvertTo-Json -Compress -Depth 5
} catch {
  [ordered]@{
    error = $_.Exception.Message
  } | ConvertTo-Json -Compress -Depth 3
  exit 1
}
            """.strip(),
            encoding="utf-8",
        )
        cls._script_path = script_path
        return script_path

    def recognize(self, image_path: Path) -> Dict[str, Any]:
        if not self.is_available():
            raise RuntimeError("Windows OCR is not available on this machine.")

        script_path = self._ensure_script_path()
        result = subprocess.run(
            ["powershell.exe", "-NoProfile", "-File", str(script_path), str(image_path)],
            capture_output=True,
            text=True,
            timeout=self.timeout_seconds,
            check=False,
        )
        output = (result.stdout or result.stderr or "").strip()
        if not output:
            raise RuntimeError("Windows OCR returned no output.")

        payload = json.loads(output)
        if result.returncode != 0:
            raise RuntimeError(str(payload.get("error") or output))

        return payload


class DocumentParser(Protocol):
    name: str

    def supports(self, file_path: Path) -> bool:
        ...

    def iter_elements(self, file_path: Path) -> Iterator[ParsedDocumentElement]:
        ...


class BaseDocumentParser:
    name = "base"
    supported_suffixes: Sequence[str] = ()

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in self.supported_suffixes

    def iter_elements(self, file_path: Path) -> Iterator[ParsedDocumentElement]:
        raise NotImplementedError

    def _make_element(
        self,
        content: str,
        element_index: int,
        metadata: Optional[Dict[str, Any]] = None,
        content_type: str = "text",
        structural_role: str = "body",
    ) -> ParsedDocumentElement:
        return ParsedDocumentElement(
            content=content,
            metadata=metadata or {},
            content_type=content_type,
            structural_role=structural_role,
            parser_name=self.name,
            element_index=element_index,
        )

    def _infer_content_type(self, text: str) -> str:
        stripped = text.strip()
        if not stripped:
            return "text"

        lines = [line.strip() for line in stripped.splitlines() if line.strip()]
        if stripped.startswith("![") or re.search(r"!\[[^\]]*\]\([^)]+\)", stripped):
            return "image_reference"
        if len(lines) >= 2 and sum(1 for line in lines if "|" in line) >= 2:
            return "table"
        return "text"

    def _infer_structural_role(self, text: str) -> str:
        stripped = " ".join(text.split())
        if not stripped:
            return "body"

        words = stripped.split()
        if len(words) <= 14:
            if stripped.isupper():
                return "heading"
            if not stripped.endswith((".", "!", "?", ";", ":")):
                title_case_words = sum(1 for word in words if word[:1].isupper())
                if title_case_words >= max(1, len(words) // 2):
                    return "heading"
        return "body"

    def _yield_text_blocks(
        self,
        text: str,
        metadata_builder,
    ) -> Iterator[ParsedDocumentElement]:
        element_index = 0
        for block_index, block in enumerate(re.split(r"\n{2,}", text), start=1):
            cleaned = block.strip()
            if not cleaned:
                continue
            yield self._make_element(
                content=cleaned,
                element_index=element_index,
                metadata=metadata_builder(block_index, cleaned),
                content_type=self._infer_content_type(cleaned),
                structural_role=self._infer_structural_role(cleaned),
            )
            element_index += 1


class PdfDocumentParser(BaseDocumentParser):
    name = "pdf"
    supported_suffixes = (".pdf",)
    line_merge_tolerance = 3.0
    paragraph_gap = 18.0
    image_suffixes = (".bmp", ".gif", ".jpeg", ".jpg", ".png", ".tif", ".tiff")

    def __init__(self):
        self.ocr_engine = WindowsOCREngine()

    def iter_elements(self, file_path: Path) -> Iterator[ParsedDocumentElement]:
        element_index = 0
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_number, page in enumerate(pdf_reader.pages, start=1):
                page_width = self._safe_float(getattr(page.mediabox, "width", None))
                page_height = self._safe_float(getattr(page.mediabox, "height", None))
                rectangles: List[Dict[str, float]] = []
                text_fragments: List[Dict[str, Any]] = []

                def visitor_rect(op: bytes, args: Any, _cm: Any, _tm: Any) -> None:
                    if op != b"re" or len(args) < 4:
                        return

                    try:
                        x, y, width, height = (float(args[index].as_numeric()) for index in range(4))
                    except Exception:
                        return

                    if width <= 1 or height <= 1:
                        return

                    rectangles.append(
                        {
                            "x": x,
                            "y": y,
                            "width": width,
                            "height": height,
                        }
                    )

                def visitor_text(text: str, cm: Any, tm: Any, _font_dict: Any, font_size: float) -> None:
                    cleaned = re.sub(r"\s+", " ", text or "").strip()
                    if not cleaned:
                        return

                    x = self._matrix_position(cm, tm, 4)
                    y = self._matrix_position(cm, tm, 5)
                    text_fragments.append(
                        {
                            "text": cleaned,
                            "x": x,
                            "y": y,
                            "font_size": float(font_size or 0.0),
                        }
                    )

                try:
                    page.extract_text(
                        visitor_operand_before=visitor_rect,
                        visitor_text=visitor_text,
                    )
                except Exception as exc:
                    logger.warning("PDF visitor extraction failed for %s page %s: %s", file_path, page_number, exc)

                lines = self._group_pdf_fragments_into_lines(text_fragments)
                if lines:
                    blocks = self._group_pdf_lines_into_blocks(
                        lines=lines,
                        rectangles=rectangles,
                        page_height=page_height,
                    )
                    for page_block_index, block in enumerate(blocks, start=1):
                        content = block["text"].strip()
                        if not content:
                            continue

                        yield self._make_element(
                            content=content,
                            element_index=element_index,
                            metadata={
                                "page_number": page_number,
                                "block_index": page_block_index,
                                "page_width": page_width,
                                "page_height": page_height,
                                "line_count": block["line_count"],
                                "y_top": block["y_top"],
                                "y_bottom": block["y_bottom"],
                                "x_min": block["x_min"],
                                "x_max": block["x_max"],
                                "font_size_avg": block["font_size_avg"],
                                "layout_mode": "visitor_layout",
                                "table_like": block["content_type"] == "table",
                            },
                            content_type=block["content_type"],
                            structural_role=block["structural_role"],
                        )
                        element_index += 1
                    continue

                page_text = page.extract_text() or ""
                for page_block_index, block in enumerate(re.split(r"\n{2,}", page_text), start=1):
                    cleaned = block.strip()
                    if not cleaned:
                        continue
                    yield self._make_element(
                        content=cleaned,
                        element_index=element_index,
                        metadata={
                            "page_number": page_number,
                            "block_index": page_block_index,
                            "page_width": page_width,
                            "page_height": page_height,
                            "layout_mode": "plain_text_fallback",
                        },
                        content_type=self._infer_content_type(cleaned),
                        structural_role=self._infer_structural_role(cleaned),
                    )
                    element_index += 1

                ocr_elements = self._extract_ocr_elements_from_page_images(
                    page=page,
                    page_number=page_number,
                    page_width=page_width,
                    page_height=page_height,
                    start_element_index=element_index,
                )
                for element in ocr_elements:
                    yield element
                element_index += len(ocr_elements)

    def _group_pdf_fragments_into_lines(self, fragments: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        line_map: Dict[float, List[Dict[str, Any]]] = defaultdict(list)
        for fragment in fragments:
            bucket = round(float(fragment["y"]) / self.line_merge_tolerance) * self.line_merge_tolerance
            line_map[bucket].append(fragment)

        lines: List[Dict[str, Any]] = []
        for y_bucket in sorted(line_map.keys(), reverse=True):
            line_fragments = sorted(line_map[y_bucket], key=lambda item: (float(item["x"]), item["text"]))
            text_parts: List[str] = []
            font_sizes: List[float] = []
            x_positions: List[float] = []
            previous_right: Optional[float] = None

            for fragment in line_fragments:
                fragment_text = str(fragment["text"]).strip()
                if not fragment_text:
                    continue

                x = float(fragment["x"])
                font_size = float(fragment["font_size"])
                if previous_right is not None:
                    gap = x - previous_right
                    if gap > max(14.0, font_size * 1.5):
                        text_parts.append("   ")
                    elif text_parts and not text_parts[-1].endswith(" "):
                        text_parts.append(" ")

                text_parts.append(fragment_text)
                font_sizes.append(font_size)
                x_positions.append(x)
                previous_right = x + max(font_size, len(fragment_text) * max(font_size, 1.0) * 0.45)

            line_text = "".join(text_parts).strip()
            if not line_text:
                continue

            lines.append(
                {
                    "text": line_text,
                    "y": float(y_bucket),
                    "x_min": min(x_positions) if x_positions else 0.0,
                    "x_max": previous_right or (max(x_positions) if x_positions else 0.0),
                    "font_size": (sum(font_sizes) / len(font_sizes)) if font_sizes else 0.0,
                }
            )

        return lines

    def _group_pdf_lines_into_blocks(
        self,
        lines: List[Dict[str, Any]],
        rectangles: List[Dict[str, float]],
        page_height: Optional[float],
    ) -> List[Dict[str, Any]]:
        if not lines:
            return []

        median_font_size = self._median([float(line["font_size"]) for line in lines if line["font_size"] > 0])
        blocks: List[Dict[str, Any]] = []
        current_block: List[Dict[str, Any]] = []
        previous_y: Optional[float] = None

        for line in lines:
            content_type = self._classify_pdf_line_content_type(line["text"], line["y"], rectangles)
            structural_role = self._classify_pdf_line_role(
                text=line["text"],
                font_size=float(line["font_size"]),
                median_font_size=median_font_size,
                y=float(line["y"]),
                page_height=page_height,
            )
            enriched_line = {
                **line,
                "content_type": content_type,
                "structural_role": structural_role,
            }

            should_split = False
            if current_block:
                gap = (previous_y - float(line["y"])) if previous_y is not None else 0.0
                if gap > self.paragraph_gap:
                    should_split = True
                elif structural_role == "heading":
                    should_split = True
                elif current_block[-1]["content_type"] != content_type and "table" in {
                    str(current_block[-1]["content_type"]),
                    content_type,
                }:
                    should_split = True

            if should_split and current_block:
                blocks.append(self._build_pdf_block(current_block))
                current_block = []

            current_block.append(enriched_line)
            previous_y = float(line["y"])

        if current_block:
            blocks.append(self._build_pdf_block(current_block))

        return blocks

    def _build_pdf_block(self, lines: List[Dict[str, Any]]) -> Dict[str, Any]:
        content_type = "table" if any(line["content_type"] == "table" for line in lines) else "text"
        non_body_roles = [str(line["structural_role"]) for line in lines if line["structural_role"] != "body"]
        structural_role = non_body_roles[0] if non_body_roles else "body"

        return {
            "text": "\n".join(str(line["text"]).strip() for line in lines if str(line["text"]).strip()),
            "content_type": content_type,
            "structural_role": structural_role,
            "line_count": len(lines),
            "y_top": max(float(line["y"]) for line in lines),
            "y_bottom": min(float(line["y"]) for line in lines),
            "x_min": min(float(line["x_min"]) for line in lines),
            "x_max": max(float(line["x_max"]) for line in lines),
            "font_size_avg": sum(float(line["font_size"]) for line in lines) / max(1, len(lines)),
        }

    def _classify_pdf_line_content_type(
        self,
        text: str,
        y: float,
        rectangles: List[Dict[str, float]],
    ) -> str:
        inferred = self._infer_content_type(text)
        if inferred != "text":
            return inferred

        if self._line_intersects_rectangle(y, rectangles):
            return "table"

        if re.search(r"\S+\s{3,}\S+", text):
            return "table"

        return "text"

    def _classify_pdf_line_role(
        self,
        text: str,
        font_size: float,
        median_font_size: float,
        y: float,
        page_height: Optional[float],
    ) -> str:
        inferred = self._infer_structural_role(text)
        if inferred == "heading":
            return inferred

        if median_font_size > 0 and font_size >= median_font_size * 1.25 and len(text.split()) <= 18:
            return "heading"

        if page_height:
            if y > page_height * 0.93 and len(text.split()) <= 10:
                return "page_header"
            if y < page_height * 0.07 and len(text.split()) <= 10:
                return "page_footer"

        return "body"

    def _line_intersects_rectangle(self, y: float, rectangles: List[Dict[str, float]]) -> bool:
        for rectangle in rectangles:
            rect_y = float(rectangle["y"])
            rect_height = float(rectangle["height"])
            if rect_y <= y <= rect_y + rect_height:
                return True
        return False

    def _matrix_position(self, cm: Any, tm: Any, index: int) -> float:
        for matrix in (tm, cm):
            try:
                value = matrix[index]
            except Exception:
                continue
            try:
                return float(value)
            except Exception:
                continue
        return 0.0

    def _median(self, values: List[float]) -> float:
        if not values:
            return 0.0

        ordered = sorted(values)
        midpoint = len(ordered) // 2
        if len(ordered) % 2 == 1:
            return ordered[midpoint]
        return (ordered[midpoint - 1] + ordered[midpoint]) / 2.0

    def _safe_float(self, value: Any) -> Optional[float]:
        if value is None:
            return None
        try:
            return float(value)
        except Exception:
            return None

    def _extract_ocr_elements_from_page_images(
        self,
        page: Any,
        page_number: int,
        page_width: Optional[float],
        page_height: Optional[float],
        start_element_index: int,
    ) -> List[ParsedDocumentElement]:
        if not self.ocr_engine.is_available():
            return []

        elements: List[ParsedDocumentElement] = []
        page_images = getattr(page, "images", []) or []
        next_element_index = start_element_index

        for image_index, image_file in enumerate(page_images, start=1):
            image_name = str(getattr(image_file, "name", f"page-{page_number}-image-{image_index}.bin"))
            image_data = getattr(image_file, "data", None)
            if not image_data:
                continue

            suffix = Path(image_name).suffix.lower()
            if suffix not in self.image_suffixes:
                continue

            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as handle:
                handle.write(image_data)
                temp_image_path = Path(handle.name)

            try:
                ocr_result = self.ocr_engine.recognize(temp_image_path)
            except Exception as exc:
                logger.warning(
                    "OCR failed for PDF page image %s on page %s: %s",
                    image_name,
                    page_number,
                    exc,
                )
                temp_image_path.unlink(missing_ok=True)
                continue

            temp_image_path.unlink(missing_ok=True)
            text = str(ocr_result.get("text") or "").strip()
            if not text:
                continue

            elements.append(
                self._make_element(
                    content=text,
                    element_index=next_element_index,
                    metadata={
                        "page_number": page_number,
                        "block_index": image_index,
                        "page_width": page_width,
                        "page_height": page_height,
                        "layout_mode": "ocr_embedded_image",
                        "ocr_source": "pdf_embedded_image",
                        "ocr_engine": "windows_media_ocr",
                        "ocr_confidence": ocr_result.get("ocr_confidence"),
                        "ocr_line_count": int(ocr_result.get("line_count") or 0),
                        "ocr_word_count": int(ocr_result.get("word_count") or 0),
                        "image_name": image_name,
                        "image_width": ocr_result.get("image_width"),
                        "image_height": ocr_result.get("image_height"),
                        "image_index": image_index,
                    },
                    content_type="image",
                    structural_role="body",
                )
            )
            next_element_index += 1

        return elements


class ImageDocumentParser(BaseDocumentParser):
    name = "image"
    supported_suffixes = (".bmp", ".gif", ".jpeg", ".jpg", ".png", ".tif", ".tiff")

    def __init__(self):
        self.ocr_engine = WindowsOCREngine()

    def iter_elements(self, file_path: Path) -> Iterator[ParsedDocumentElement]:
        if not self.ocr_engine.is_available():
            raise ValueError("Windows OCR is not available for image document parsing.")

        ocr_result = self.ocr_engine.recognize(file_path)
        text = str(ocr_result.get("text") or "").strip()
        if not text:
            return

        yield self._make_element(
            content=text,
            element_index=0,
            metadata={
                "layout_mode": "ocr_image_file",
                "ocr_source": "image_file",
                "ocr_engine": "windows_media_ocr",
                "ocr_confidence": ocr_result.get("ocr_confidence"),
                "ocr_line_count": int(ocr_result.get("line_count") or 0),
                "ocr_word_count": int(ocr_result.get("word_count") or 0),
                "image_width": ocr_result.get("image_width"),
                "image_height": ocr_result.get("image_height"),
                "image_name": file_path.name,
            },
            content_type="image",
            structural_role="body",
        )


class SpreadsheetDocumentParser(BaseDocumentParser):
    name = "spreadsheet"
    supported_suffixes = (".xlsx", ".csv", ".tsv")
    block_row_limit = 25

    def iter_elements(self, file_path: Path) -> Iterator[ParsedDocumentElement]:
        suffix = file_path.suffix.lower()
        if suffix == ".xlsx":
            yield from self._iter_xlsx_elements(file_path)
            return

        delimiter = "\t" if suffix == ".tsv" else ","
        yield from self._iter_delimited_elements(file_path, delimiter=delimiter)

    def _iter_delimited_elements(self, file_path: Path, delimiter: str) -> Iterator[ParsedDocumentElement]:
        with file_path.open("r", encoding="utf-8-sig", newline="") as handle:
            rows = list(csv.reader(handle, delimiter=delimiter))

        normalized_rows = self._normalize_tabular_rows(rows)
        yield from self._iter_table_blocks(
            table_rows=normalized_rows,
            element_prefix=file_path.stem,
            metadata_base={
                "sheet_name": file_path.stem,
                "sheet_index": 1,
                "layout_mode": "delimited_table",
                "delimiter": delimiter,
            },
        )

    def _iter_xlsx_elements(self, file_path: Path) -> Iterator[ParsedDocumentElement]:
        namespace = {
            "main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
            "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
        }

        with zipfile.ZipFile(file_path) as archive:
            workbook_root = ET.fromstring(archive.read("xl/workbook.xml"))
            rels_root = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
            relationships = {
                rel.attrib.get("Id"): rel.attrib.get("Target", "")
                for rel in rels_root.findall("rel:Relationship", namespace)
            }
            shared_strings = self._load_shared_strings(archive, namespace)

            for sheet_index, sheet in enumerate(workbook_root.findall("main:sheets/main:sheet", namespace), start=1):
                relationship_id = sheet.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
                target = relationships.get(relationship_id)
                if not target:
                    continue

                target_path = target if target.startswith("xl/") else f"xl/{target}"
                if target_path not in archive.namelist():
                    continue

                sheet_name = sheet.attrib.get("name", f"Sheet{sheet_index}")
                sheet_rows = self._load_sheet_rows(
                    archive=archive,
                    sheet_path=target_path,
                    shared_strings=shared_strings,
                    namespace=namespace,
                )
                yield from self._iter_table_blocks(
                    table_rows=sheet_rows,
                    element_prefix=sheet_name,
                    metadata_base={
                        "workbook_name": file_path.name,
                        "sheet_name": sheet_name,
                        "sheet_index": sheet_index,
                        "layout_mode": "xlsx_table",
                    },
                )

    def _load_shared_strings(self, archive: zipfile.ZipFile, namespace: Dict[str, str]) -> List[str]:
        if "xl/sharedStrings.xml" not in archive.namelist():
            return []

        root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
        values: List[str] = []
        for shared_string in root.findall("main:si", namespace):
            text_parts = [
                node.text or ""
                for node in shared_string.findall(".//main:t", namespace)
            ]
            values.append("".join(text_parts))
        return values

    def _load_sheet_rows(
        self,
        archive: zipfile.ZipFile,
        sheet_path: str,
        shared_strings: List[str],
        namespace: Dict[str, str],
    ) -> List[List[str]]:
        root = ET.fromstring(archive.read(sheet_path))
        rows: List[List[str]] = []

        for row in root.findall(".//main:sheetData/main:row", namespace):
            values_by_index: Dict[int, str] = {}
            for cell in row.findall("main:c", namespace):
                cell_ref = cell.attrib.get("r", "")
                column_index = self._column_index_from_reference(cell_ref)
                if column_index is None:
                    continue

                value = self._extract_cell_value(cell, shared_strings, namespace)
                values_by_index[column_index] = value

            if not values_by_index:
                continue

            max_index = max(values_by_index.keys())
            normalized_row = [values_by_index.get(index, "") for index in range(max_index + 1)]
            rows.append(normalized_row)

        return self._normalize_tabular_rows(rows)

    def _extract_cell_value(
        self,
        cell: ET.Element,
        shared_strings: List[str],
        namespace: Dict[str, str],
    ) -> str:
        cell_type = cell.attrib.get("t", "")
        inline_text = "".join(
            node.text or ""
            for node in cell.findall("main:is//main:t", namespace)
        )
        if inline_text:
            return inline_text.strip()

        value_node = cell.find("main:v", namespace)
        value = (value_node.text or "").strip() if value_node is not None else ""
        if not value:
            return ""

        if cell_type == "s":
            try:
                return shared_strings[int(value)]
            except Exception:
                return value

        if cell_type == "b":
            return "TRUE" if value == "1" else "FALSE"

        return value

    def _column_index_from_reference(self, cell_reference: str) -> Optional[int]:
        column_part = "".join(character for character in cell_reference if character.isalpha()).upper()
        if not column_part:
            return None

        index = 0
        for character in column_part:
            index = index * 26 + (ord(character) - ord("A") + 1)
        return index - 1

    def _normalize_tabular_rows(self, rows: List[List[str]]) -> List[List[str]]:
        normalized = [
            [str(cell).strip() for cell in row]
            for row in rows
            if any(str(cell).strip() for cell in row)
        ]
        if not normalized:
            return []

        width = max(len(row) for row in normalized)
        return [row + [""] * (width - len(row)) for row in normalized]

    def _iter_table_blocks(
        self,
        table_rows: List[List[str]],
        element_prefix: str,
        metadata_base: Dict[str, Any],
    ) -> Iterator[ParsedDocumentElement]:
        if not table_rows:
            return

        headers = table_rows[0]
        data_rows = table_rows[1:] if len(table_rows) > 1 else []
        has_header = any(header.strip() for header in headers) and len(data_rows) > 0
        body_rows = data_rows if has_header else table_rows

        if not body_rows and has_header:
            body_rows = [headers]

        for block_index, start in enumerate(range(0, len(body_rows), self.block_row_limit), start=1):
            block_rows = body_rows[start:start + self.block_row_limit]
            row_start = start + (2 if has_header else 1)
            row_end = row_start + len(block_rows) - 1
            table_text = self._render_table_block(
                sheet_label=element_prefix,
                headers=headers if has_header else None,
                rows=block_rows,
                row_start=row_start,
            )
            metadata = {
                **metadata_base,
                "block_index": block_index,
                "row_start": row_start,
                "row_end": row_end,
                "column_count": max((len(row) for row in block_rows), default=len(headers)),
                "has_header_row": has_header,
            }
            yield self._make_element(
                content=table_text,
                element_index=block_index - 1,
                metadata=metadata,
                content_type="table",
                structural_role="body",
            )

    def _render_table_block(
        self,
        sheet_label: str,
        headers: Optional[List[str]],
        rows: List[List[str]],
        row_start: int,
    ) -> str:
        lines = [f"Sheet: {sheet_label}"]
        if headers:
            header_line = " | ".join(header.strip() or f"Column {index + 1}" for index, header in enumerate(headers))
            lines.append(f"Headers: {header_line}")

        for offset, row in enumerate(rows):
            row_number = row_start + offset
            row_values = [cell.strip() or "-" for cell in row]
            lines.append(f"Row {row_number}: {' | '.join(row_values)}")

        return "\n".join(lines).strip()


class DocxDocumentParser(BaseDocumentParser):
    name = "docx"
    supported_suffixes = (".docx",)

    def iter_elements(self, file_path: Path) -> Iterator[ParsedDocumentElement]:
        doc = Document(file_path)
        element_index = 0
        for paragraph_index, paragraph in enumerate(doc.paragraphs, start=1):
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = (paragraph.style.name or "").strip() if paragraph.style else ""
            style_name_lower = style_name.lower()
            structural_role = (
                "heading"
                if any(token in style_name_lower for token in ("heading", "title", "subtitle"))
                else self._infer_structural_role(text)
            )
            yield self._make_element(
                content=text,
                element_index=element_index,
                metadata={
                    "paragraph_index": paragraph_index,
                    "style_name": style_name,
                },
                content_type=self._infer_content_type(text),
                structural_role=structural_role,
            )
            element_index += 1


class MarkdownDocumentParser(BaseDocumentParser):
    name = "markdown"
    supported_suffixes = (".md",)

    def iter_elements(self, file_path: Path) -> Iterator[ParsedDocumentElement]:
        text = file_path.read_text(encoding='utf-8')
        element_index = 0
        for block_index, block in enumerate(re.split(r"\n{2,}", text), start=1):
            cleaned = block.strip()
            if not cleaned:
                continue

            heading_match = re.match(r"^(#{1,6})\s+(.*)$", cleaned)
            if heading_match:
                content = heading_match.group(2).strip()
                structural_role = "heading"
                metadata = {
                    "markdown_block_index": block_index,
                    "heading_level": len(heading_match.group(1)),
                }
            else:
                content = cleaned
                structural_role = self._infer_structural_role(cleaned)
                metadata = {
                    "markdown_block_index": block_index,
                }

            yield self._make_element(
                content=content,
                element_index=element_index,
                metadata=metadata,
                content_type=self._infer_content_type(cleaned),
                structural_role=structural_role,
            )
            element_index += 1


class TextDocumentParser(BaseDocumentParser):
    name = "text"
    supported_suffixes = (".txt",)

    def iter_elements(self, file_path: Path) -> Iterator[ParsedDocumentElement]:
        text = file_path.read_text(encoding='utf-8')
        yield from self._yield_text_blocks(
            text,
            metadata_builder=lambda block_index, _: {
                "text_block_index": block_index,
            },
        )

class SmartChunker:
    """Intelligent document chunking with semantic boundaries, from your script."""
    def __init__(self, chunk_size: int = 512, overlap: int = 128, parsers: Optional[List[DocumentParser]] = None):
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.parsers = parsers or self._default_parsers()
        
    def chunk_text(self, text: str, metadata: Dict) -> List[DocumentChunk]:
        """Chunk text with modality-aware rules and overlap where it is useful."""
        cleaned = text.strip()
        if not cleaned:
            return []

        content_type = str(metadata.get("content_type") or "text")
        structural_role = str(metadata.get("structural_role") or "body")

        if structural_role in {"heading", "title", "subtitle"}:
            return [self._make_chunk(cleaned, metadata, chunk_strategy="structural_atomic")]

        if content_type == "table":
            return self._chunk_table_content(cleaned, metadata)

        if content_type in {"image", "figure", "image_reference"}:
            return self._chunk_line_preserving_content(
                cleaned,
                metadata,
                chunk_strategy="visual_lines",
                overlap_units=1,
            )

        return self._chunk_prose_content(cleaned, metadata)

    def _chunk_prose_content(self, text: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Chunk narrative text on sentence boundaries with token overlap."""
        sentences = self._split_into_sentences(text)
        chunks: List[DocumentChunk] = []
        current_chunk_words: List[str] = []
        
        for sentence in sentences:
            sentence_words = sentence.split()
            if not sentence_words:
                continue

            if len(sentence_words) > self.chunk_size:
                if current_chunk_words:
                    chunks.append(
                        self._make_chunk(
                            " ".join(current_chunk_words),
                            metadata,
                            chunk_strategy="semantic_sentences",
                        )
                    )
                    overlap_word_count = min(len(current_chunk_words), self.overlap)
                    current_chunk_words = current_chunk_words[-overlap_word_count:] if overlap_word_count else []

                step = max(1, self.chunk_size - self.overlap)
                for start in range(0, len(sentence_words), step):
                    window_words = sentence_words[start:start + self.chunk_size]
                    if not window_words:
                        continue
                    chunks.append(
                        self._make_chunk(
                            " ".join(window_words),
                            metadata,
                            chunk_strategy="semantic_sentences",
                        )
                    )

                overlap_word_count = min(len(sentence_words), self.overlap)
                current_chunk_words = sentence_words[-overlap_word_count:] if overlap_word_count else []
                continue

            if len(current_chunk_words) + len(sentence_words) > self.chunk_size and current_chunk_words:
                chunks.append(
                    self._make_chunk(
                        " ".join(current_chunk_words),
                        metadata,
                        chunk_strategy="semantic_sentences",
                    )
                )
                
                overlap_word_count = min(len(current_chunk_words), self.overlap)
                current_chunk_words = current_chunk_words[-overlap_word_count:]

            current_chunk_words.extend(sentence_words)

        if current_chunk_words:
            chunks.append(
                self._make_chunk(
                    " ".join(current_chunk_words),
                    metadata,
                    chunk_strategy="semantic_sentences",
                )
            )
        
        return chunks

    def _chunk_table_content(self, text: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Chunk table-derived content without destroying row boundaries."""
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if not lines:
            return []

        prelude: List[str] = []
        row_entries: List[Tuple[Optional[int], str]] = []
        row_pattern = re.compile(r"^Row\s+(\d+):", re.IGNORECASE)

        for line in lines:
            match = row_pattern.match(line)
            if match:
                row_entries.append((int(match.group(1)), line))
            else:
                prelude.append(line)

        if not row_entries:
            return self._chunk_line_preserving_content(
                text,
                metadata,
                chunk_strategy="table_lines",
                overlap_units=0,
            )

        chunks: List[DocumentChunk] = []
        current_rows: List[Tuple[Optional[int], str]] = []
        prefix_word_count = self._count_words("\n".join(prelude))

        for row_number, row_text in row_entries:
            projected_words = prefix_word_count + self._count_words("\n".join(line for _, line in current_rows))
            projected_words += self._count_words(row_text)

            if current_rows and projected_words > self.chunk_size:
                chunks.append(self._render_table_chunk(prelude, current_rows, metadata))
                current_rows = []

            current_rows.append((row_number, row_text))

        if current_rows:
            chunks.append(self._render_table_chunk(prelude, current_rows, metadata))

        return chunks

    def _render_table_chunk(
        self,
        prelude: List[str],
        row_entries: List[Tuple[Optional[int], str]],
        metadata: Dict[str, Any],
    ) -> DocumentChunk:
        content_lines = [*prelude, *(row_text for _, row_text in row_entries)]
        chunk_metadata = dict(metadata)
        row_numbers = [row_number for row_number, _ in row_entries if row_number is not None]
        if row_numbers:
            chunk_metadata["table_row_start"] = min(row_numbers)
            chunk_metadata["table_row_end"] = max(row_numbers)
            chunk_metadata["table_row_count"] = len(row_numbers)
        return self._make_chunk(
            "\n".join(content_lines).strip(),
            chunk_metadata,
            chunk_strategy="table_rows",
        )

    def _chunk_line_preserving_content(
        self,
        text: str,
        metadata: Dict[str, Any],
        chunk_strategy: str,
        overlap_units: int = 0,
    ) -> List[DocumentChunk]:
        """Chunk OCR or layout-driven text while preserving line grouping."""
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if not lines:
            return []

        chunks: List[DocumentChunk] = []
        current_lines: List[str] = []
        current_word_count = 0

        for line in lines:
            line_word_count = self._count_words(line)
            if current_lines and current_word_count + line_word_count > self.chunk_size:
                chunks.append(
                    self._make_chunk(
                        "\n".join(current_lines).strip(),
                        metadata,
                        chunk_strategy=chunk_strategy,
                    )
                )
                overlap_lines = current_lines[-overlap_units:] if overlap_units else []
                current_lines = list(overlap_lines)
                current_word_count = self._count_words("\n".join(current_lines))

            if line_word_count > self.chunk_size:
                if current_lines:
                    chunks.append(
                        self._make_chunk(
                            "\n".join(current_lines).strip(),
                            metadata,
                            chunk_strategy=chunk_strategy,
                        )
                    )
                    current_lines = []
                    current_word_count = 0

                chunks.extend(self._split_oversized_line(line, metadata, chunk_strategy))
                continue

            current_lines.append(line)
            current_word_count += line_word_count

        if current_lines:
            chunks.append(
                self._make_chunk(
                    "\n".join(current_lines).strip(),
                    metadata,
                    chunk_strategy=chunk_strategy,
                )
            )

        return chunks

    def _split_oversized_line(
        self,
        line: str,
        metadata: Dict[str, Any],
        chunk_strategy: str,
    ) -> List[DocumentChunk]:
        words = line.split()
        if not words:
            return []

        chunks: List[DocumentChunk] = []
        step = max(1, self.chunk_size - self.overlap)
        for start in range(0, len(words), step):
            window_words = words[start:start + self.chunk_size]
            if not window_words:
                continue
            chunks.append(
                self._make_chunk(
                    " ".join(window_words),
                    metadata,
                    chunk_strategy=chunk_strategy,
                )
            )
        return chunks

    def _make_chunk(
        self,
        content: str,
        metadata: Dict[str, Any],
        chunk_strategy: str,
    ) -> DocumentChunk:
        return DocumentChunk(
            content=content,
            metadata={**metadata, "chunk_strategy": chunk_strategy},
        )

    def _count_words(self, text: str) -> int:
        return len(text.split())

    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences."""
        normalized = text.replace('\r\n', '\n').replace('\r', '\n')
        segments = []

        for block in re.split(r'\n{2,}', normalized):
            block = block.strip()
            if not block:
                continue

            for line in block.split('\n'):
                cleaned_line = re.sub(r'^[#>*\-\+\s]+', '', line).strip()
                if not cleaned_line:
                    continue

                parts = re.split(r'(?<=[.!?])\s+', cleaned_line)
                for part in parts:
                    part = re.sub(r'\s+', ' ', part).strip()
                    if part:
                        segments.append(part)

        return segments

    def process_document(self, file_path: str) -> List[DocumentChunk]:
        """Process a supported document file into chunks."""
        chunks: List[DocumentChunk] = []
        for batch in self.iter_document_chunk_batches(file_path=file_path, batch_size=256):
            chunks.extend(batch)
        return chunks

    def parse_document(self, file_path: str) -> ParsedDocument:
        file_p = Path(file_path)
        return ParsedDocument(
            source=str(file_p),
            filename=file_p.name,
            elements=list(self._iter_parsed_elements(file_p)),
        )

    def iter_document_chunk_batches(self, file_path: str, batch_size: int = 32) -> Iterator[List[DocumentChunk]]:
        """Yield chunk batches without requiring the full document to stay in memory."""
        file_p = Path(file_path)
        if not file_p.exists():
            raise FileNotFoundError(f"File not found: {file_p}")

        base_metadata = {
            'source': str(file_p),
            'filename': file_p.name,
            'processed_at': datetime.now().isoformat()
        }

        pending_chunks: List[DocumentChunk] = []
        next_chunk_index = 0

        try:
            for element in self._iter_parsed_elements(file_p):
                if not element.content.strip():
                    continue

                metadata = {
                    **base_metadata,
                    **element.metadata,
                    "content_type": element.content_type,
                    "structural_role": element.structural_role,
                    "parser_name": element.parser_name,
                    "element_index": element.element_index,
                }
                for chunk in self.chunk_text(element.content, metadata):
                    chunk.metadata['chunk_index'] = next_chunk_index
                    next_chunk_index += 1
                    pending_chunks.append(chunk)

                    if len(pending_chunks) >= batch_size:
                        yield pending_chunks
                        pending_chunks = []
        except Exception as e:
            logger.error(f"Failed to process {file_path}: {e}")
            return

        if pending_chunks:
            yield pending_chunks

    def _extract_pdf_text(self, file_path: Path) -> str:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
        return text

    def _extract_docx_text(self, file_path: Path) -> str:
        doc = Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])

    def _extract_markdown_text(self, file_path: Path) -> str:
        return file_path.read_text(encoding='utf-8')

    def _extract_text_file(self, file_path: Path) -> str:
        return file_path.read_text(encoding='utf-8')

    def _default_parsers(self) -> List[DocumentParser]:
        return [
            PdfDocumentParser(),
            ImageDocumentParser(),
            SpreadsheetDocumentParser(),
            DocxDocumentParser(),
            MarkdownDocumentParser(),
            TextDocumentParser(),
        ]

    def _iter_parsed_elements(self, file_path: Path) -> Iterator[ParsedDocumentElement]:
        parser = self._get_parser(file_path)
        yield from parser.iter_elements(file_path)

    def _get_parser(self, file_path: Path) -> DocumentParser:
        for parser in self.parsers:
            if parser.supports(file_path):
                return parser
        raise ValueError(f"Unsupported file type: {file_path.suffix}")

    def _iter_document_segments(self, file_path: Path) -> Iterator[Tuple[str, Dict[str, Any]]]:
        for element in self._iter_parsed_elements(file_path):
            yield element.content, {
                **element.metadata,
                "content_type": element.content_type,
                "structural_role": element.structural_role,
                "parser_name": element.parser_name,
                "element_index": element.element_index,
            }


class EnhancedChunker(SmartChunker):
    """Enhanced chunking with entity-aware splitting, dynamic sizing, topic boundary detection, cross-chunk references, and enhanced metadata"""
    def __init__(self, chunk_size: int = 512, overlap: int = 128, use_ner: bool = True, min_chunk_size: int = 256, max_chunk_size: int = 1024, detect_topics: bool = True):
        super().__init__(chunk_size, overlap)
        self.use_ner = use_ner
        self.min_chunk_size = min_chunk_size
        self.max_chunk_size = max_chunk_size
        self.detect_topics = detect_topics
        try:
            import spacy
            self.nlp = spacy.load("en_core_web_sm", disable=['parser'])
        except Exception as e:
            logger.warning(f"Could not load spaCy model for NER: {e}. Falling back to basic chunking.")
            self.use_ner = False
            self.detect_topics = False
            
    def chunk_text(self, text: str, metadata: Dict) -> List[DocumentChunk]:
        """Chunk text with entity-aware boundaries, dynamic sizing, topic detection, cross-chunk references, enhanced metadata, and overlap"""
        if not self.use_ner:
            return super().chunk_text(text, metadata)
            
        # Process text with spaCy for entity recognition
        doc = self.nlp(text)
        sentences = list(doc.sents)
        chunks = []
        current_chunk_tokens = []
        current_entities = set()
        current_entity_types = {}
        current_complexity = 0
        current_topic = None
        topic_counter = 0
        document_id = metadata.get('source', 'doc') + '_' + str(hash(text) if len(text) > 0 else 'empty')
        content_types = set()
        
        for i, sent in enumerate(sentences):
            sent_tokens = [token.text for token in sent]
            sent_entities = {ent.text for ent in sent.ents}
            sent_entity_types = {ent.text: ent.label_ for ent in sent.ents}
            # Estimate complexity based on entity density and sentence length
            sent_complexity = len(sent_entities) * 2 + len(sent_tokens) / 10
            
            # Detect content type (rudimentary)
            if any(token.text.lower() in ['chapter', 'section', 'article'] for token in sent):
                content_types.add('structural')
            if any(ent.label_ in ['PERSON', 'ORG'] for ent in sent.ents):
                content_types.add('narrative')
            if any(ent.label_ in ['DATE', 'TIME', 'MONEY', 'QUANTITY'] for ent in sent.ents):
                content_types.add('factual')
            
            # Detect potential topic shift (rudimentary approach based on entities and keywords)
            topic_shift = False
            if self.detect_topics and i > 0 and current_topic is not None:
                prev_sent = sentences[i-1]
                prev_entities = {ent.text for ent in prev_sent.ents}
                common_entities = len(sent_entities.intersection(prev_entities))
                # If few common entities, possible topic shift
                if common_entities < len(sent_entities) * 0.3 and common_entities < len(prev_entities) * 0.3:
                    topic_shift = True
            
            # Check if adding this sentence would exceed dynamic chunk size or if topic shift detected
            current_size = len(current_chunk_tokens)
            dynamic_threshold = min(self.max_chunk_size, max(self.min_chunk_size, self.chunk_size - int(current_complexity)))
            if (current_size + len(sent_tokens) > dynamic_threshold or topic_shift) and current_chunk_tokens:
                # Avoid splitting entities if possible
                if not any(ent in current_entities for ent in sent_entities):
                    chunk_content = " ".join(current_chunk_tokens)
                    chunk_id = f"{document_id}_chunk_{len(chunks)}"
                    chunk_metadata = {
                        **metadata,
                        'chunk_index': len(chunks),
                        'chunk_id': chunk_id,
                        'entities': list(current_entities),
                        'entity_types': current_entity_types,
                        'complexity': current_complexity / max(1, current_size) if current_size > 0 else 0,
                        'topic_id': current_topic if current_topic is not None else f"topic_{topic_counter}",
                        'content_types': list(content_types),
                        'prev_chunk_id': f"{document_id}_chunk_{len(chunks)-1}" if len(chunks) > 0 else None,
                        'next_chunk_id': None  # Will be updated in the next iteration
                    }
                    # Update the next_chunk_id of the previous chunk if it exists
                    if len(chunks) > 0:
                        chunks[-1].metadata['next_chunk_id'] = chunk_id
                    
                    chunks.append(DocumentChunk(content=chunk_content, metadata=chunk_metadata))
                    
                    # Start new chunk with overlap
                    overlap_token_count = min(len(current_chunk_tokens), self.overlap)
                    current_chunk_tokens = current_chunk_tokens[-overlap_token_count:]
                    current_entities = set()
                    current_entity_types = {}
                    current_complexity = 0
                    content_types = set()
                    for token in current_chunk_tokens:
                        for ent in doc.ents:
                            if token in ent.text:
                                current_entities.add(ent.text)
                                if ent.text in sent_entity_types:
                                    current_entity_types[ent.text] = sent_entity_types[ent.text]
                    if topic_shift:
                        topic_counter += 1
                        current_topic = f"topic_{topic_counter}"
                    else:
                        current_topic = chunk_metadata['topic_id']
            
            current_chunk_tokens.extend(sent_tokens)
            current_entities.update(sent_entities)
            current_entity_types.update(sent_entity_types)
            current_complexity += sent_complexity
            if current_topic is None:
                current_topic = f"topic_{topic_counter}"
        
        if current_chunk_tokens:
            chunk_content = " ".join(current_chunk_tokens)
            current_size = len(current_chunk_tokens)
            chunk_id = f"{document_id}_chunk_{len(chunks)}"
            chunk_metadata = {
                **metadata,
                'chunk_index': len(chunks),
                'chunk_id': chunk_id,
                'entities': list(current_entities),
                'entity_types': current_entity_types,
                'complexity': current_complexity / max(1, current_size) if current_size > 0 else 0,
                'topic_id': current_topic if current_topic is not None else f"topic_{topic_counter}",
                'content_types': list(content_types),
                'prev_chunk_id': f"{document_id}_chunk_{len(chunks)-1}" if len(chunks) > 0 else None,
                'next_chunk_id': None
            }
            # Update the next_chunk_id of the previous chunk if it exists
            if len(chunks) > 0:
                chunks[-1].metadata['next_chunk_id'] = chunk_id
            
            chunks.append(DocumentChunk(content=chunk_content, metadata=chunk_metadata))
            
        return chunks

    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences"""
        text = re.sub(r'\n+', ' ', text)
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s and s.strip()]

    def process_document(self, file_path: str) -> List[DocumentChunk]:
        """Process a document file into chunks"""
        file_p = Path(file_path)
        if not file_p.exists():
            raise FileNotFoundError(f"File not found: {file_p}")
        
        text = ""
        try:
            if file_p.suffix.lower() == '.pdf':
                text = self._extract_pdf_text(file_p)
            elif file_p.suffix.lower() == '.docx':
                text = self._extract_docx_text(file_p)
            elif file_p.suffix.lower() == '.md':
                text = self._extract_markdown_text(file_p)
            elif file_p.suffix.lower() == '.txt':
                text = self._extract_text_file(file_p)
            else:
                raise ValueError(f"Unsupported file type: {file_p.suffix}")
        except Exception as e:
            logger.error(f"Failed to process {file_path}: {e}")
            return []

        metadata = {
            'source': str(file_p),
            'filename': file_p.name,
            'processed_at': datetime.now().isoformat()
        }
        return self.chunk_text(text, metadata)
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
        return text
    
    def _extract_docx_text(self, file_path: Path) -> str:
        doc = Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    def _extract_markdown_text(self, file_path: Path) -> str:
        with open(file_path, 'r', encoding='utf-8') as file:
            html = markdown.markdown(file.read())
            return re.sub(r'<[^>]+>', '', html)
    
    def _extract_text_file(self, file_path: Path) -> str:
        return file_path.read_text(encoding='utf-8')
